"""Clean executive DOCX export for ATS Career Builder.

Drop-in replacement for core/docx_export.py.
Public entry point is unchanged:  build_docx(resume_text: str) -> bytes
so api/main.py needs NO changes.

Fixes:
- No spell-check (red) or grammar (blue) squiggles.
- Normal-weight body text (no more bold+italic everywhere).
- Tight, even spacing - no large blank gaps.
- Company left / dates right via a real tab stop -> no tables, no dotted gridlines.
- Skills sanitizer: drops junk tokens (IT, IP, HR, BI, LI, AND, NIT, CGPA, GATE,
  SAF, HP, EEM, DXC, GRS, SGS, USA, ET, CIO, GCC, VAT ...) so the skills line
  stays professional even when the backend feeds garbage.
- De-duplicates a header/contact line that gets repeated inside the summary.
- Fixes letter-spaced headings like "E X E C U T I V E  E X P E R T I S E".
- Clickable LinkedIn hyperlink.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
NAVY_HEX = "1F3A5F"
INK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)
RIGHT_TAB_IN = 7.1
HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

_SECTION_HEADS = {
    "PROFILE SUMMARY", "PROFESSIONAL SUMMARY", "PROFILE", "SUMMARY",
    "KEY ACHIEVEMENTS", "ENTERPRISE ACCOMPLISHMENTS", "SELECTED CAREER HIGHLIGHTS",
    "CAREER HIGHLIGHTS", "CORE COMPETENCIES", "CORE COMPETENCIES & TECHNOLOGIES",
    "CORE LEADERSHIP & TECHNICAL EXPERTISE", "EXECUTIVE EXPERTISE",
    "TECHNICAL SKILLS", "TECHNICAL ENVIRONMENT", "KEY SKILLS",
    "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE",
    "EARLIER PROFESSIONAL EXPERIENCE",
    "EDUCATION & CERTIFICATIONS", "EDUCATION AND CERTIFICATION",
    "EDUCATION AND CERTIFICATIONS", "EDUCATION", "CERTIFICATIONS", "RECOGNITION",
}
_INLINE_SECTIONS = {
    "CORE COMPETENCIES", "CORE COMPETENCIES & TECHNOLOGIES",
    "CORE LEADERSHIP & TECHNICAL EXPERTISE", "EXECUTIVE EXPERTISE",
    "TECHNICAL SKILLS", "TECHNICAL ENVIRONMENT", "KEY SKILLS",
}
_EXPERIENCE_SECTIONS = {
    "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE",
    "EARLIER PROFESSIONAL EXPERIENCE",
}

_DATE_RANGE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{4})"
    r"\s*(?:-|\u2013|\u2014|to)\s*"
    r"((?:present|current)|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{4})",
    re.I,
)

_SKILL_STOPLIST = {
    "it", "ip", "hr", "bi", "li", "and", "nit", "cgpa", "gate", "saf", "hp",
    "dxc", "eem", "grs", "sgs", "usa", "et", "cio", "gcc", "vat", "qa", "pmo",
    "kt", "rca", "sla", "the", "a", "an", "of", "to", "in", "on", "for", "with",
}
_SKILL_KEEP = {"sap", "aws", "gcp", "ci/cd", "sdlc", "api", "rag", "ai", "j2ee", "pmp"}


def _is_junk_skill(token: str) -> bool:
    t = token.strip().lower()
    if not t or t in _SKILL_STOPLIST:
        return True
    if len(t) <= 2 and t not in {"ai"}:
        return True
    raw = token.strip()
    if raw.isupper() and len(raw) <= 4 and t not in _SKILL_KEEP:
        return True
    return False


def _clean_skills(items):
    seen, out = set(), []
    for it in items:
        it = it.strip(" \t|\u2022\u25AA-")
        if not it or _is_junk_skill(it):
            continue
        k = it.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(it)
    return out


def _fix_letter_spacing(text: str) -> str:
    s = text.strip()
    letters = [c for c in s if c != " "]
    if len(letters) >= 4 and s.count(" ") >= len(letters) - 1 and s.replace(" ", "").isalpha():
        words = re.split(r"\s{2,}", s)
        return " ".join(w.replace(" ", "") for w in words)
    return text


def _run(p, text, *, size=10.3, bold=False, italic=False, color=INK, caps=False, font="Calibri"):
    r = p.add_run(text.upper() if caps else text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), font); rf.set(qn("w:hAnsi"), font)
    if rpr.find(qn("w:noProof")) is None:
        rpr.append(OxmlElement("w:noProof"))
    return r


def _sp(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _border(p, color=NAVY_HEX, size="6"):
    pPr = p._p.get_or_add_pPr()
    b = pPr.find(qn("w:pBdr"))
    if b is None:
        b = OxmlElement("w:pBdr"); pPr.append(b)
    e = OxmlElement("w:bottom")
    for k, v in (("val", "single"), ("sz", size), ("space", "3"), ("color", color)):
        e.set(qn("w:" + k), v)
    b.append(e)


def _keep_next(p):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))


def _hyperlink(p, text, url):
    rid = p.part.relate_to(url, HYPERLINK_REL, is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rid)
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), NAVY_HEX); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), "Calibri"); rf.set(qn("w:hAnsi"), "Calibri"); rpr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20"); rpr.append(sz)
    rpr.append(OxmlElement("w:noProof")); r.append(rpr)
    t = OxmlElement("w:t"); t.text = text; r.append(t); h.append(r); p._p.append(h)


def _disable_proofing(doc):
    st = doc.settings.element
    for tag in ("w:hideSpellingErrors", "w:hideGrammaticalErrors"):
        st.append(OxmlElement(tag))


def _heading(doc, text):
    text = _fix_letter_spacing(text)
    p = doc.add_paragraph(); _sp(p, before=7, after=2, line=1.0); _keep_next(p)
    r = _run(p, text, size=11.5, bold=True, color=NAVY, caps=True)
    rpr = r._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "16"); rpr.append(sp)
    _border(p)


def _paragraph(doc, text):
    p = doc.add_paragraph(); _sp(p, after=3, line=1.12)
    _run(p, text, size=10.1, color=INK)


def _bullet(doc, text):
    p = doc.add_paragraph(); _sp(p, after=2, line=1.08)
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    _run(p, "\u2022  ", size=10, bold=True, color=NAVY)
    _run(p, text, size=10.2, color=INK)


def _inline_list(doc, items):
    items = _clean_skills(items)
    if not items:
        return
    p = doc.add_paragraph(); _sp(p, after=3, line=1.22)
    for i, it in enumerate(items):
        if i:
            _run(p, "   |   ", size=10, bold=True, color=NAVY)
        _run(p, it, size=10.2, color=INK)


def _experience_row(doc, company, dates, title):
    p = doc.add_paragraph(); _sp(p, before=5, after=0, line=1.0); _keep_next(p)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(RIGHT_TAB_IN), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    _run(p, company, size=10.8, bold=True, color=NAVY)
    if dates:
        _run(p, "\t"); _run(p, dates, size=9.6, bold=True, color=NAVY)
    if title:
        pt = doc.add_paragraph(); _sp(pt, after=2, line=1.0); _keep_next(pt)
        _run(pt, title, size=10, italic=True, color=GRAY)


def _split_lines(text):
    return [l.rstrip() for l in (text or "").replace("\r\n", "\n").split("\n")]


def _is_heading(line):
    return _fix_letter_spacing(line).strip().upper().rstrip(":") in _SECTION_HEADS


def build_docx(resume_text: str) -> bytes:
    doc = Document()
    _disable_proofing(doc)

    for s in doc.sections:
        s.top_margin = Inches(0.45); s.bottom_margin = Inches(0.45)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    lines = _split_lines(resume_text)
    nonempty = [i for i, l in enumerate(lines) if l.strip()]

    name = headline = contact = ""
    consumed = 0
    if nonempty:
        name = lines[nonempty[0]].strip(); consumed = nonempty[0] + 1
        if len(nonempty) > 1 and not _is_heading(lines[nonempty[1]]):
            headline = lines[nonempty[1]].strip(); consumed = nonempty[1] + 1
        if len(nonempty) > 2 and not _is_heading(lines[nonempty[2]]) and any(
                c in lines[nonempty[2]] for c in ("@", "|", "+")):
            contact = lines[nonempty[2]].strip(); consumed = nonempty[2] + 1

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(p, after=1)
    _run(p, name or "YOUR NAME", size=22, bold=True, color=NAVY, caps=True)
    if headline:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(p, after=2)
        _run(p, headline, size=11, bold=True, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(p, after=3)
    if contact:
        m = re.search(r"(https?://)?(www\.)?linkedin\.com/\S+", contact, re.I)
        if m:
            before = contact[:m.start()].rstrip(" |\u2022")
            _run(p, (before + "  |  ") if before else "", size=9.4, color=INK)
            url = m.group(0)
            if not url.lower().startswith("http"):
                url = "https://" + url
            _hyperlink(p, m.group(0), url)
        else:
            _run(p, contact, size=9.4, color=INK)
    _border(p, NAVY_HEX, "8")

    sig = re.sub(r"\s+", " ", (name + " " + headline)).strip().lower()[:40]

    current = None
    i = consumed
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1; continue

        if _is_heading(line):
            current = _fix_letter_spacing(line).upper().rstrip(":")
            _heading(doc, current)
            i += 1; continue

        low = re.sub(r"\s+", " ", line).lower()
        if sig and sig in low and ("@" in line or "|" in line or current in {"PROFESSIONAL SUMMARY", "PROFILE SUMMARY"}):
            i += 1; continue

        if current in _INLINE_SECTIONS:
            parts = re.split(r"\s*[|\u2022\u25AA]\s*", line)
            _inline_list(doc, [x for x in parts if x.strip()])
            i += 1; continue

        if current in _EXPERIENCE_SECTIONS:
            m = _DATE_RANGE.search(line)
            if m:
                dates = m.group(0)
                company = line[:m.start()].strip(" ,-\u2013\u2014|\t")
                title = ""
                if i + 1 < len(lines):
                    nxt = lines[i + 1].strip()
                    if (nxt and nxt[:2] not in ("- ", "* ", "\u2022 ", "\u25AA ")
                            and not _DATE_RANGE.search(nxt) and not _is_heading(nxt)):
                        title = nxt; i += 1
                _experience_row(doc, company or line, dates, title)
                i += 1; continue

        if line[:2] in ("- ", "* ", "\u2022 ", "\u25AA ") or line[:1] in ("\u2022", "\u25AA"):
            _bullet(doc, re.sub(r"^[\-\*\u2022\u25AA]\s*", "", line))
            i += 1; continue

        _paragraph(doc, _fix_letter_spacing(line))
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
