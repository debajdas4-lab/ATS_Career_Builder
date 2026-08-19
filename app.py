from __future__ import annotations

import html
import os
import re
from collections import Counter
from io import BytesIO
from typing import Any

import httpx
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

st.set_page_config(
    page_title="ATS Career Builder V3",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

API_URL = os.getenv("RESUME_API_URL", "http://127.0.0.1:8000").strip().rstrip("/")

# -----------------------------------------------------------------------------
# UI
# -----------------------------------------------------------------------------
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
:root{--navy:#13293f;--navy2:#081a31;--gold:#c2a14d;--blue:#315be8;--violet:#5b45ee;--ink:#17233d;--muted:#6d7890;--line:#e3e8f2}
html,body,[class*="css"]{font-family:Inter,sans-serif}.stApp{background:linear-gradient(180deg,#fff,#f5f8fd);color:var(--ink)}
.block-container{max-width:1450px;padding-top:0!important;padding-bottom:3rem}header[data-testid="stHeader"]{height:2.7rem;background:rgba(255,255,255,.96)}
section[data-testid="stSidebar"]{background:linear-gradient(180deg,var(--navy2),#17365d)}section[data-testid="stSidebar"] *{color:#e8edff}.brand{padding:.75rem .6rem 1.2rem;border-bottom:1px solid rgba(255,255,255,.12);margin-bottom:1rem}.brand h2{margin:0 0 .35rem;font-size:1.35rem;font-weight:800}.brand span{color:#e7c76e}.brand p{margin:0;color:#b8c6dd;font-size:.82rem}.nav-label{color:#8496b8!important;font-size:.7rem;font-weight:800;letter-spacing:.16em;margin:1rem .4rem .45rem}.nav,.nav-active{padding:.75rem .8rem;border-radius:12px;margin:.16rem 0;font-weight:600;font-size:.9rem}.nav-active{background:linear-gradient(90deg,#315be8,#6550e8);box-shadow:0 7px 22px rgba(68,81,235,.3)}
.hero{padding:1.12rem 2.5rem;min-height:142px;border-radius:0 0 14px 14px;background:radial-gradient(circle at 78% 47%,rgba(69,128,255,.55),transparent 18%),linear-gradient(115deg,#081a31 0,#12335c 52%,#27217c 78%,#081a31 100%);color:#fff;box-shadow:0 8px 30px rgba(8,23,51,.18)}.hero-kicker{color:#44d6e8;font-size:.92rem;font-weight:600}.hero h1{font-size:2.4rem;line-height:1.02;margin:.25rem 0 .3rem;letter-spacing:-.04em}.hero p{font-size:.94rem;font-weight:600;color:#dfe8ff;margin:.2rem 0 .7rem}.pills{display:flex;gap:.5rem;flex-wrap:wrap}.pill{padding:.32rem .6rem;border:1px solid rgba(121,170,255,.24);background:rgba(15,37,85,.72);border-radius:999px;font-size:.72rem;color:#cddafd}.pill b{color:#47d9eb}
.workspace-kicker{text-align:center;color:#375ee8;font-size:.9rem;font-weight:800;letter-spacing:.16em;margin-top:1.05rem}.workspace-title{font-size:1.7rem;font-weight:800;letter-spacing:-.04em;margin:.5rem 0 .25rem}.workspace-copy{font-size:.95rem;color:var(--muted);line-height:1.58;margin-bottom:1.1rem}.section-title{font-size:1.07rem;font-weight:800;margin:1.25rem 0 .3rem}.or-line{display:flex;align-items:center;gap:.8rem;margin:.7rem 0;color:#4f5f7d;font-size:.76rem;font-weight:800;letter-spacing:.12em}.or-line:before,.or-line:after{content:"";height:1px;background:#dce3ef;flex:1}
.stTabs [data-baseweb="tab-list"]{display:flex;flex-wrap:wrap;gap:.42rem;background:#eef2f8;padding:.45rem;border:1px solid #e1e7f0;border-radius:14px}.stTabs [data-baseweb="tab"]{height:2.65rem;border-radius:10px;padding:0 1rem;color:#56627a;font-size:.84rem;font-weight:700}.stTabs [aria-selected="true"]{background:linear-gradient(90deg,#17365d,#5147c7)!important;color:#fff!important}.stTabs [data-baseweb="tab-panel"]{background:#fff;border:1px solid #e3e8f2;border-radius:0 0 14px 14px;padding:1.25rem 1.35rem;margin-top:.5rem}.score-card{background:#fff;border:1px solid #e2e7f0;border-radius:14px;padding:1.1rem 1.2rem;box-shadow:0 5px 18px rgba(23,35,61,.05)}.score-title{font-size:.8rem;font-weight:800;color:#66738a;letter-spacing:.08em}.score-value{font-size:2rem;font-weight:800;color:#17365d;margin:.25rem 0}.score-delta{font-size:.85rem;font-weight:700;color:#17785f}.bar{height:8px;background:#e9edf5;border-radius:99px;overflow:hidden}.bar>span{display:block;height:100%;background:linear-gradient(90deg,#315be8,#5b45ee);border-radius:99px}.chip{display:inline-block;padding:.32rem .58rem;border-radius:999px;margin:.15rem;background:#edf2ff;color:#304fc9;font-size:.74rem;font-weight:700;border:1px solid #d7e0ff}.chip.missing{background:#fff3ec;color:#a3521e;border-color:#ffd9c4}.resume{background:#fff;border:1px solid var(--line);border-radius:14px;padding:1.65rem 1.8rem;line-height:1.72;max-width:1000px}.resume h3{color:#13293f;border-bottom:1px solid #c2a14d;padding-bottom:.3rem;margin-top:1.25rem}.fit-note{background:#f2f6ff;border-left:4px solid #315be8;padding:.9rem 1rem;border-radius:8px;color:#435069;line-height:1.55}
div[data-testid="stButton"]>button[kind="primary"],.stButton>button[kind="primary"]{background:linear-gradient(90deg,#315be8,#4165ee,#5b45ee)!important;color:#fff!important;border:0!important;border-radius:11px!important;font-weight:800!important;box-shadow:0 8px 20px rgba(49,91,232,.28)!important}
</style>
""",
    unsafe_allow_html=True,
)

# -----------------------------------------------------------------------------
# Input extraction
# -----------------------------------------------------------------------------
def extract_uploaded_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(data)).pages)
    if name.endswith(".docx"):
        document = Document(BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
        return text
    return data.decode("utf-8", errors="ignore")

# -----------------------------------------------------------------------------
# Before/after fit dashboard
# -----------------------------------------------------------------------------
STOPWORDS = {
    "the", "and", "for", "with", "from", "that", "this", "will", "you", "your",
    "our", "are", "have", "has", "into", "across", "within", "while", "their",
    "they", "who", "its", "but", "not", "all", "using", "through", "strong",
    "ability", "experience", "including", "minimum", "preferred", "demonstrated",
    "proven", "excellent", "successful", "multiple", "closely", "ensure", "role",
    "team", "teams", "work", "working", "lead", "drive", "driving", "manage",
}
PHRASES = [
    "technical program management", "cross-functional", "program governance",
    "roadmap alignment", "stakeholder management", "risk management",
    "operational excellence", "software development lifecycle", "cloud technologies",
    "system architecture", "program execution", "success metrics", "senior leadership",
    "agile", "devops", "ci/cd", "sap s/4hana", "vendor management",
]
ALIASES = {
    "technical program management": ["technical program manager", "program management"],
    "cross-functional": ["cross functional", "global teams"],
    "roadmap alignment": ["roadmap", "milestones"],
    "risk management": ["risk mitigation", "raid", "risk tracking"],
    "software development lifecycle": ["sdlc"],
    "cloud technologies": ["cloud", "paas", "azure", "aws"],
    "stakeholder management": ["stakeholder engagement", "stakeholder leadership"],
    "program governance": ["governance framework", "governance", "kpi governance"],
}


def clean_text(value: str) -> str:
    text = BeautifulSoup(value or "", "html.parser").get_text(" ")
    return re.sub(r"\s+", " ", text.lower().replace("s/4 hana", "s/4hana")).strip()


def extract_requirements(jd: str) -> list[str]:
    text = clean_text(jd)
    requirements = [phrase for phrase in PHRASES if phrase in text]
    counts = Counter(
        word for word in re.findall(r"[a-z][a-z0-9/+.-]{2,}", text)
        if word not in STOPWORDS
    )
    requirements.extend(word for word, count in counts.most_common(35) if count >= 2)
    return list(dict.fromkeys(requirements))[:45]


def score_resume(resume: str, jd: str) -> dict:
    requirements = extract_requirements(jd)
    text = clean_text(resume)
    matched = [
        req for req in requirements
        if any(candidate in text for candidate in [req] + ALIASES.get(req, []))
    ]
    return {
        "score": round(100 * len(matched) / len(requirements)) if requirements else 0,
        "matched": matched,
        "missing": [req for req in requirements if req not in matched],
    }

# -----------------------------------------------------------------------------
# Source completeness and Complete Career Guide fallback
# -----------------------------------------------------------------------------
EMPLOYERS = {
    "mindtree": (
        "Mindtree Limited | Technical Lead | Bangalore, India | Jan 2007 - Sep 2010",
        "Led enterprise programs for Volvo, SGS, and Travelport while introducing Scrum and Test-Driven Development across global teams.",
    ),
    "hewlett packard": (
        "Hewlett Packard Pvt. Ltd. | Senior Software Engineer | Bangalore, India | Aug 2004 - Dec 2006",
        "Engineered enterprise procurement solutions and contributed to Web Presentation and Shared Services architecture modernization.",
    ),
    "netkraft": (
        "Netkraft and Metafusion | Software Engineer / Engineer | Bangalore, India | Feb 2002 - Aug 2004",
        "Developed enterprise information and material-management applications for global customers across design, testing, performance optimization, integration, and production delivery.",
    ),
}
SECTIONS = {
    "PROFESSIONAL SUMMARY", "SELECTED CAREER HIGHLIGHTS",
    "CORE LEADERSHIP & TECHNICAL EXPERTISE", "PROFESSIONAL EXPERIENCE",
    "EARLIER PROFESSIONAL EXPERIENCE", "EDUCATION & CERTIFICATIONS",
}


def correct_resume(generated: str, source: str) -> tuple[str, list[str]]:
    lines = [re.sub(r"\s+", " ", line).strip() for line in generated.splitlines() if line.strip()]
    notes: list[str] = []
    joined = "\n".join(lines).lower()
    missing = []
    for key, entry in EMPLOYERS.items():
        if key in source.lower() and key not in joined:
            missing.append(entry)
            notes.append(f"Restored employer: {entry[0].split('|')[0].strip()}")
    if missing:
        index = next(
            (i for i, line in enumerate(lines) if line.upper().rstrip(":") == "EDUCATION & CERTIFICATIONS"),
            len(lines),
        )
        if "EARLIER PROFESSIONAL EXPERIENCE" not in [line.upper().rstrip(":") for line in lines]:
            lines.insert(index, "EARLIER PROFESSIONAL EXPERIENCE")
            index += 1
        for role, bullet in missing:
            lines.insert(index, role)
            lines.insert(index + 1, "- " + bullet)
            index += 2
    return "\n".join(lines), notes


def _list(value: Any) -> list[str]:
    if isinstance(value, (list, tuple, set)):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _unique(values: list[str], limit: int = 25) -> list[str]:
    output: list[str] = []
    seen: set[str] = set()
    for value in values:
        item = re.sub(r"\s+", " ", value).strip(" -|•")
        if item and item.lower() not in seen:
            seen.add(item.lower())
            output.append(item)
        if len(output) >= limit:
            break
    return output


def ensure_complete_mode_outputs(result: dict, source_resume: str, jd: str) -> dict:
    """Fill blank Complete Career Guide modules without extra model calls."""
    recruiter = result.get("recruiter_intelligence") or {}
    brand = result.get("executive_branding") or {}
    keywords = _unique(
        _list(recruiter.get("recruiter_keywords"))
        + _list(recruiter.get("ats_keywords"))
        + _list(recruiter.get("leadership_keywords"))
        + _list(recruiter.get("technical_keywords")),
        30,
    )
    headline = brand.get("executive_headline") or (
        "Senior Technical Program Manager | Enterprise Transformation | "
        "Program Governance | SAP S/4HANA | Agile Delivery"
    )
    value = brand.get("value_proposition") or (
        "Technical program management leader focused on enterprise transformation, "
        "governance, cross-functional delivery, risk management, and operational excellence."
    )

    linkedin = result.get("linkedin_optimization")
    if not isinstance(linkedin, dict) or not any(linkedin.values()):
        result["linkedin_optimization"] = {
            "headline": headline,
            "about_or_summary": value,
            "featured_capabilities": keywords[:12],
            "skills": keywords[:20],
            "keywords": keywords[:25],
            "recommended_changes": [
                "Use the target-role headline while keeping the current title accurate.",
                "Use only source-supported outcomes and quantified achievements.",
                "Align the Skills section with supported recruiter keywords.",
                "Keep transformation, governance, stakeholder leadership, SAP S/4HANA, and Agile delivery easy to find when supported.",
            ],
            "generation_note": "Source-based fallback used because the backend returned no LinkedIn content.",
        }

    interview = result.get("interview_kit")
    if not isinstance(interview, dict) or not any(interview.values()):
        bullets = [
            line[2:].strip() for line in (result.get("optimized_resume") or source_resume).splitlines()
            if line.strip().startswith(("- ", "• ", "▪ ", "* "))
        ]
        gaps = _unique(
            _list((result.get("keyword_gap") or {}).get("missing"))
            + _list((result.get("ats_analysis") or {}).get("missing_keywords")),
            10,
        )
        result["interview_kit"] = {
            "resume_questions": [
                "Walk through the most complex transformation program and explain scope, governance, dependencies, and outcomes.",
                "Describe how risks, issues, and cross-team dependencies were identified, escalated, and resolved.",
                "Explain one example of aligning business, engineering, infrastructure, QA, operations, and vendors.",
            ],
            "leadership_questions": [
                "How do you establish governance without slowing engineering delivery?",
                "How do you resolve conflict between technical constraints, roadmap priorities, and regulatory deadlines?",
                "How do you communicate delivery health, financial exposure, and decisions to senior leaders?",
            ],
            "technical_or_domain_questions": [
                "How do you collaborate with architects and engineering leads on technical trade-offs?",
                "How do release planning, testing, production readiness, and hypercare fit into your operating model?",
                "How have SAP S/4HANA, cloud modernization, Agile, DevOps, or automation been applied in supported programs?",
            ],
            "gap_questions": [
                f"The role emphasizes '{gap}'. What related evidence can you discuss without overstating direct experience?"
                for gap in gaps[:6]
            ],
            "star_story_blueprints": [
                {
                    "source_evidence": bullet,
                    "situation": "Describe the business or delivery context.",
                    "task": "Clarify the accountable outcome, constraints, and stakeholders.",
                    "action": "Explain the governance, technical coordination, risk controls, and decisions led.",
                    "result": "Use only outcomes and metrics explicitly supported by the source resume.",
                }
                for bullet in bullets[:5]
            ],
            "generation_note": "Source-based fallback used because the backend returned no Interview Kit content.",
        }
    return result

# -----------------------------------------------------------------------------
# DOCX export
# -----------------------------------------------------------------------------
NAVY = "224F79"
BLACK = "222222"
GRAY = "646464"
BLUE = "5B9BD5"


def set_run(run, size=9.2, bold=False, color=BLACK, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def bottom_border(paragraph, color=NAVY, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def keep_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def keep_lines(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def add_hyperlink(paragraph, text, url):
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(
        12 if title in {"SELECTED CAREER HIGHLIGHTS", "CORE LEADERSHIP & TECHNICAL EXPERTISE"} else 9
    )
    paragraph.paragraph_format.space_after = Pt(5)
    keep_next(paragraph)
    set_run(paragraph.add_run(title.upper()), 9.8, True, NAVY)
    bottom_border(paragraph)


def add_bullet(document, text, compact=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.14)
    paragraph.paragraph_format.space_after = Pt(3.2 if not compact else 2.3)
    paragraph.paragraph_format.line_spacing = 1.16 if not compact else 1.10
    keep_lines(paragraph)
    set_run(paragraph.add_run("• "), 8, False, BLUE)
    set_run(paragraph.add_run(text), 9.2 if not compact else 8.9)


def normalize_role(line: str) -> str:
    parts = [part.strip() for part in line.split("|")]
    if len(parts) < 4:
        return line
    date, location = parts[-1], parts[-2]
    first = parts[0]
    second = " | ".join(parts[1:-2])
    markers = ("inc", "limited", "systems", "wipro", "mindtree", "hewlett", "netkraft", "metafusion")
    company, title = (first, second) if any(marker in first.lower() for marker in markers) else (second, first)
    return f"{company} | {title} | {location} | {date}"


def add_role(document, line):
    parts = [part.strip() for part in normalize_role(line).split("|")]
    company = parts[0]
    title = parts[1] if len(parts) > 1 else ""
    location = parts[2] if len(parts) > 2 else ""
    dates = parts[3] if len(parts) > 3 else ""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5.5)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.95), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    keep_next(paragraph)
    set_run(paragraph.add_run(company), 9.9, True, BLACK)
    set_run(paragraph.add_run("\t" + dates), 8.3, True, NAVY)
    sub = document.add_paragraph()
    sub.paragraph_format.space_after = Pt(2.8)
    sub.paragraph_format.tab_stops.add_tab_stop(Inches(6.95), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    keep_next(sub)
    set_run(sub.add_run(title), 8.6, False, NAVY, True)
    set_run(sub.add_run("\t" + location), 8.1, False, GRAY, True)


def is_role_line(line: str) -> bool:
    return "|" in line and bool(re.search(r"\b(?:19|20)\d{2}\b|\bPresent\b", line, re.I))


def create_docx(text: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)
    section.footer_distance = Inches(0.22)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(3.0)
    normal.paragraph_format.line_spacing = 1.16
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Deba Jyoti Das | Technical Program Management | Page "), 7.2, False, GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    header_index = 0
    current_section = ""
    for line in lines:
        upper = line.upper().rstrip(":")
        if upper in SECTIONS:
            current_section = upper
            section_heading(document, upper)
            continue
        if upper in {"NAME", "HEADLINE", "POSITIONING LINE", "CONTACT"}:
            continue
        if header_index < 4:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt({0: 1.5, 1: 2.5, 2: 3, 3: 9}[header_index])
            if header_index == 0:
                set_run(paragraph.add_run(line.upper()), 20, True, NAVY)
            elif header_index == 1:
                set_run(paragraph.add_run(line.upper()), 10.6, True, BLACK)
            elif header_index == 2:
                set_run(paragraph.add_run(line), 9, False, GRAY)
            else:
                linkedin = "linkedin.com/in/debajyoti-das"
                if linkedin in line.lower():
                    before = line[: line.lower().index(linkedin)].rstrip(" •|")
                    set_run(paragraph.add_run(before + " • "), 8.15, False, BLACK)
                    add_hyperlink(paragraph, linkedin, "https://linkedin.com/in/debajyoti-das")
                else:
                    set_run(paragraph.add_run(line), 8.15, False, BLACK)
                bottom_border(paragraph, NAVY, "5")
            header_index += 1
            continue
        if current_section in {"PROFESSIONAL EXPERIENCE", "EARLIER PROFESSIONAL EXPERIENCE"} and is_role_line(line):
            add_role(document, line)
            continue
        if line.startswith(("- ", "• ", "▪ ", "* ")):
            add_bullet(document, line[2:].strip(), current_section == "EARLIER PROFESSIONAL EXPERIENCE")
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4.2)
        paragraph.paragraph_format.line_spacing = 1.16
        keep_lines(paragraph)
        set_run(paragraph.add_run(line), 9.2)

    output = BytesIO()
    document.save(output)
    return output.getvalue()

# -----------------------------------------------------------------------------
# Rendering helpers
# -----------------------------------------------------------------------------
def plain_text(value: Any) -> str | list[str]:
    text = str(value).strip()
    if "<" in text and ">" in text:
        soup = BeautifulSoup(text, "html.parser")
        items = [item.get_text(" ", strip=True) for item in soup.find_all("li")]
        return items or soup.get_text(" ", strip=True)
    return text


def render_structured(value: Any, empty="No content generated."):
    if value is None or value == "" or value == [] or value == {}:
        st.info(empty)
        return
    if isinstance(value, dict):
        for key, item in value.items():
            st.markdown(f"#### {html.escape(str(key).replace('_', ' ').title())}")
            render_structured(item, empty)
        return
    if isinstance(value, (list, tuple, set)):
        for item in value:
            if isinstance(item, (dict, list, tuple, set)):
                render_structured(item, empty)
            else:
                display = plain_text(item)
                if isinstance(display, list):
                    for entry in display:
                        st.markdown("- " + html.escape(entry))
                elif display:
                    st.markdown("- " + html.escape(display))
        return
    display = plain_text(value)
    if isinstance(display, list):
        for entry in display:
            st.markdown("- " + html.escape(entry))
    elif display:
        st.markdown(html.escape(display))


def chips(items, missing=False):
    css = "chip missing" if missing else "chip"
    return "".join(f'<span class="{css}">{html.escape(str(item))}</span>' for item in items or [])


def preview(text):
    safe = html.escape(text or "")
    for name in SECTIONS:
        safe = re.sub(rf"(?m)^{re.escape(name)}:?$", f"<h3>{name.title()}</h3>", safe, flags=re.I)
    return f'<div class="resume">{safe.replace(chr(10), "<br>")}</div>'


def score_card(title, value, delta=None):
    delta_html = f'<div class="score-delta">{html.escape(delta)}</div>' if delta else ""
    st.markdown(
        f'<div class="score-card"><div class="score-title">{html.escape(title)}</div>'
        f'<div class="score-value">{value}/100</div><div class="bar"><span style="width:{max(0,min(100,value))}%"></span></div>{delta_html}</div>',
        unsafe_allow_html=True,
    )

# -----------------------------------------------------------------------------
# App
# -----------------------------------------------------------------------------
with st.sidebar:
    st.markdown(
        '<div class="brand"><h2>ATS <span>Career Builder</span></h2><p>AI Career Intelligence & Application Copilot</p></div>'
        '<div class="nav-label">CORE TOOLS</div><div class="nav-active">🎯 &nbsp; Analyze Job</div>'
        '<div class="nav">📄 &nbsp; Resume</div><div class="nav">in &nbsp; LinkedIn</div>'
        '<div class="nav">⌁ &nbsp; Naukri</div><div class="nav">◉ &nbsp; Interview Kit</div>'
        '<div class="nav">⌁ &nbsp; Career Roadmap</div><div class="nav">◌ &nbsp; Research</div>',
        unsafe_allow_html=True,
    )

st.markdown(
    '<div class="hero"><div class="hero-kicker">Welcome to</div><h1>ATS Career Builder</h1>'
    '<p>AI-Powered Career Intelligence. Evidence-First Results.</p><div class="pills">'
    '<span class="pill"><b>✓</b> Before/After Fit</span><span class="pill"><b>✓</b> Recruiter Intelligence</span>'
    '<span class="pill"><b>✓</b> Executive Branding</span><span class="pill"><b>✓</b> Complete Career Guide</span>'
    '</div></div><div class="workspace-kicker">CAREER INTELLIGENCE WORKSPACE</div>'
    '<div class="workspace-title">Elevate Your Career. Realize Your Potential.</div>'
    '<div class="workspace-copy">Generate a complete source-validated executive resume and career guide.</div>',
    unsafe_allow_html=True,
)

st.markdown('<div class="section-title">Upload Your Resume <span style="color:#e04848">*</span></div>', unsafe_allow_html=True)
resume_file = st.file_uploader("Resume", type=["pdf", "docx", "txt", "md"], label_visibility="collapsed")
st.markdown('<div class="section-title">Job Description Source</div>', unsafe_allow_html=True)
job_url = st.text_input("Job description URL", placeholder="Optional public LinkedIn, Naukri or company careers URL")
st.markdown('<div class="or-line">OR, PASTE JOB DESC. BELOW</div>', unsafe_allow_html=True)
job_description = st.text_area("Job description", height=180, placeholder="Paste the complete job description...")
st.markdown('<div class="section-title">Analysis Mode</div>', unsafe_allow_html=True)
analysis_mode = st.radio("Mode", ["Fast Resume", "Complete Career Guide"], horizontal=True, label_visibility="collapsed")

with st.expander("Optional company, market and profile inputs"):
    left, right = st.columns(2)
    with left:
        company_url = st.text_input("Company URL")
        linkedin_profile = st.text_area("LinkedIn Profile Data", height=100)
        naukri_profile = st.text_area("Naukri Profile Data", height=100)
    with right:
        leadership_url = st.text_input("Leadership / IR URL")
        market_query = st.text_input("Market Research Query")
        market_urls = st.text_area("Additional Public URLs", height=100)

run_disabled = resume_file is None or (len(job_description.strip()) < 80 and not job_url.strip())

if st.button("◉ Analyze Job & Build Career Guide", type="primary", use_container_width=True, disabled=run_disabled):
    status = st.status("Building your career intelligence guide", expanded=True)
    try:
        source = extract_uploaded_text(resume_file)
        response = httpx.post(
            f"{API_URL}/v3/career-guide",
            files={"resume": (resume_file.name, resume_file.getvalue(), resume_file.type or "application/octet-stream")},
            data={
                "analysis_mode": analysis_mode,
                "job_description": job_description,
                "job_url": job_url,
                "company_url": company_url,
                "leadership_url": leadership_url,
                "market_urls": market_urls,
                "market_query": market_query,
                "linkedin_profile": linkedin_profile,
                "naukri_profile": naukri_profile,
            },
            timeout=httpx.Timeout(connect=10, read=None, write=120, pool=10),
        )
        response.raise_for_status()
        result = response.json()
        if analysis_mode == "Complete Career Guide":
            result = ensure_complete_mode_outputs(result, source, job_description)
        corrected, notes = correct_resume(result.get("optimized_resume", ""), source)
        result["optimized_resume"] = corrected
        result["source_correction_notes"] = notes
        result["before_fit"] = score_resume(source, job_description)
        result["after_fit"] = score_resume(corrected, job_description)
        st.session_state["career_result"] = result
        status.update(label="Career guide ready", state="complete", expanded=False)
    except httpx.HTTPStatusError as exc:
        try:
            detail = exc.response.json().get("detail", exc.response.text)
        except Exception:
            detail = exc.response.text
        status.update(label="Analysis failed", state="error")
        st.error(f"API error {exc.response.status_code}: {detail}")
    except Exception as exc:
        status.update(label="Analysis failed", state="error")
        st.error(f"{type(exc).__name__}: {exc}")

result = st.session_state.get("career_result")
if result:
    st.divider()
    before = result.get("before_fit", {})
    after = result.get("after_fit", {})
    delta = after.get("score", 0) - before.get("score", 0)
    st.markdown("## Resume Upgrade Dashboard")
    col1, col2, col3 = st.columns(3)
    with col1:
        score_card("CURRENT RESUME FIT", before.get("score", 0))
    with col2:
        score_card("UPGRADED RESUME FIT", after.get("score", 0), f"{delta:+d} points vs. current")
    with col3:
        gain = len(after.get("matched", [])) - len(before.get("matched", []))
        score_card("ADDITIONAL MATCHED SIGNALS", max(0, gain), "Same JD and scoring method")
    st.markdown(
        '<div class="fit-note"><b>ATS Fit is an internal evidence-coverage score:</b> '
        'it compares resume evidence with extracted JD requirements. It is not a hiring probability.</div>',
        unsafe_allow_html=True,
    )

    tabs = st.tabs([
        "Fit & Gaps", "Recruiter Intelligence", "Executive Branding", "Executive Resume",
        "LinkedIn", "Naukri", "Interview Kit", "Career Roadmap", "Research",
    ])
    with tabs[0]:
        newly = [item for item in after.get("matched", []) if item not in before.get("matched", [])]
        st.markdown("### Newly covered after upgrade")
        st.markdown(chips(newly), unsafe_allow_html=True)
        st.markdown("### Remaining gaps")
        st.markdown(chips(after.get("missing", []), True), unsafe_allow_html=True)
    with tabs[1]:
        render_structured(result.get("recruiter_intelligence", {}))
    with tabs[2]:
        render_structured(result.get("executive_branding", {}))
    with tabs[3]:
        text = result.get("optimized_resume", "")
        if text:
            st.markdown(preview(text), unsafe_allow_html=True)
            notes = result.get("source_correction_notes", [])
            if notes:
                with st.expander("Source-evidence corrections applied"):
                    for note in notes:
                        st.write("- " + note)
            st.download_button(
                "Download Corrected DOCX",
                create_docx(text),
                "executive-ats-resume-corrected.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.download_button("Download TXT", text, "executive-ats-resume-corrected.txt", "text/plain")
    with tabs[4]:
        render_structured(result.get("linkedin_optimization", {}), "Select Complete Career Guide.")
    with tabs[5]:
        render_structured(result.get("naukri_optimization", {}), "Select Complete Career Guide.")
    with tabs[6]:
        render_structured(result.get("interview_kit", {}), "Select Complete Career Guide.")
    with tabs[7]:
        render_structured(result.get("career_roadmap", {}), "Select Complete Career Guide.")
    with tabs[8]:
        render_structured(result.get("research", {}), "Select Complete Career Guide and add public URLs.")
