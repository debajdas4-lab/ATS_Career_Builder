"""Resume + job-description ingestion (files and public URLs)."""
from __future__ import annotations

import io

import httpx
from bs4 import BeautifulSoup

from .config import MAX_JD_CHARS, REQUEST_TIMEOUT
from .utils import normalize_ws


def extract_resume_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _from_pdf(content)
    if name.endswith(".docx"):
        return _from_docx(content)
    # txt / md / anything else -> best-effort decode
    return normalize_ws(content.decode("utf-8", errors="ignore"))


def _from_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return normalize_ws("\n".join(pages))


def _from_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return normalize_ws("\n".join(parts))


def fetch_job_description(url: str) -> str:
    """Best-effort extraction of a public job posting. Returns '' on failure."""
    url = (url or "").strip()
    if not url:
        return ""
    try:
        headers = {"User-Agent": "Mozilla/5.0 (compatible; ATSCareerBuilder/1.0)"}
        resp = httpx.get(url, headers=headers, timeout=REQUEST_TIMEOUT, follow_redirects=True)
        resp.raise_for_status()
    except Exception:
        return ""
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    text = normalize_ws(soup.get_text("\n"))
    return text[:MAX_JD_CHARS]
